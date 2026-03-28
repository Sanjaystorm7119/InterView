import json
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session

from app.database import get_db
from app.models.user import User
from app.schemas.interview import (
    QuestionGenerateRequest,
    FeedbackGenerateRequest,
    CompanySummaryRequest,
)
from app.services.ai_service import (
    get_openrouter_client,
    sanitize,
    QUESTIONS_PROMPT,
    FEEDBACK_PROMPT,
    COMPANY_SUMMARY_PROMPT,
    DOCUMENT_EXTRACT_PROMPT,
    RESUME_EXTRACT_PROMPT,
)
from app.utils.deps import get_current_user
from app.config import get_settings

settings = get_settings()
router = APIRouter(prefix="/api/ai", tags=["ai"])

MAX_FILE_SIZE = 2 * 1024 * 1024  # 2 MB


def _clean_json_response(raw: str) -> dict:
    cleaned = raw.strip()
    if cleaned.startswith("```json"):
        cleaned = cleaned[7:]
    if cleaned.startswith("```"):
        cleaned = cleaned[3:]
    if cleaned.endswith("```"):
        cleaned = cleaned[:-3]
    cleaned = cleaned.strip()
    return json.loads(cleaned)


@router.post("/generate-questions")
def generate_questions(
    body: QuestionGenerateRequest,
    current_user: User = Depends(get_current_user),
):
    prompt = QUESTIONS_PROMPT.format(
        duration=sanitize(body.duration, 20),
        jobTitle=sanitize(body.jobPosition, 200),
        jobDescription=sanitize(body.jobDescription, 5000),
        type=sanitize(body.type, 100),
        companyDetails=sanitize(body.companyDetails, 2000),
    )

    client = get_openrouter_client()
    completion = client.chat.completions.create(
        model=settings.GEMINI_FLASH_LITE,
        messages=[{"role": "user", "content": prompt}],
    )

    if not completion.choices or not completion.choices[0].message:
        raise HTTPException(status_code=500, detail="Invalid response from AI model")

    return {"content": completion.choices[0].message.content}


@router.post("/generate-feedback")
def generate_feedback(
    body: FeedbackGenerateRequest,
    current_user: User = Depends(get_current_user),
):
    company_prefix = ""
    if body.companyDetails:
        company_prefix = f"Company Details:\n{sanitize(body.companyDetails, 2000)}\n\n"

    prompt = company_prefix + FEEDBACK_PROMPT.format(
        conversation=body.conversation[:50000]
    )

    client = get_openrouter_client()
    completion = client.chat.completions.create(
        model=settings.GEMINI_FLASH_25,
        messages=[{"role": "user", "content": prompt}],
    )

    if not completion.choices or not completion.choices[0].message:
        raise HTTPException(status_code=500, detail="Invalid response from AI model")

    return {
        "content": completion.choices[0].message.content,
        "call_id": body.call_id,
        "interview_id": body.interview_id,
        "user_email": current_user.email,
    }


@router.post("/company-summary")
def company_summary(
    body: CompanySummaryRequest,
    current_user: User = Depends(get_current_user),
):
    prompt = COMPANY_SUMMARY_PROMPT.format(
        jobPosition=sanitize(body.jobPosition, 200),
        jobDescription=sanitize(body.jobDescription, 5000),
        companyDetails=sanitize(body.companyDetails, 2000),
    )

    client = get_openrouter_client()
    completion = client.chat.completions.create(
        model=settings.GEMINI_FLASH_LITE,
        messages=[{"role": "user", "content": prompt}],
    )

    if not completion.choices or not completion.choices[0].message:
        raise HTTPException(status_code=500, detail="Invalid response from AI model")

    summary = completion.choices[0].message.content.strip()
    return {"summary": summary}


@router.post("/parse-document")
async def parse_document(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
):
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="File size exceeds the 2 MB limit")

    filename = (file.filename or "").lower()
    is_pdf = file.content_type == "application/pdf" or filename.endswith(".pdf")
    is_docx = (
        file.content_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or filename.endswith(".docx")
        or filename.endswith(".doc")
    )

    if not is_pdf and not is_docx:
        raise HTTPException(
            status_code=400, detail="Only PDF and DOCX files are supported"
        )

    if is_pdf:
        import base64

        b64 = base64.b64encode(content).decode("utf-8")
        message_content = [
            {
                "type": "file",
                "file": {
                    "filename": file.filename or "document.pdf",
                    "file_data": f"data:application/pdf;base64,{b64}",
                },
            },
            {"type": "text", "text": DOCUMENT_EXTRACT_PROMPT},
        ]
    else:
        from docx import Document
        import io

        doc = Document(io.BytesIO(content))
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        if not doc_text.strip():
            raise HTTPException(
                status_code=422, detail="Could not extract text from the document"
            )
        message_content = (
            f"Document content:\n\n{doc_text}\n\n---\n\n{DOCUMENT_EXTRACT_PROMPT}"
        )

    client = get_openrouter_client()
    completion = client.chat.completions.create(
        model=settings.GEMMA_3N,
        messages=[{"role": "user", "content": message_content}],
    )

    raw = completion.choices[0].message.content or ""
    try:
        parsed = _clean_json_response(raw)
    except json.JSONDecodeError:
        raise HTTPException(
            status_code=500,
            detail="Failed to parse model response as JSON",
        )

    return {
        "company_name": parsed.get("company_name", ""),
        "company_details": parsed.get("company_details", ""),
        "job_position": parsed.get("job_position", ""),
        "job_description": parsed.get("job_description", ""),
    }


@router.post("/parse-resume")
async def parse_resume(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
):
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="File size exceeds the 2 MB limit")

    filename = (file.filename or "").lower()
    is_pdf = file.content_type == "application/pdf" or filename.endswith(".pdf")
    is_docx = (
        file.content_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or filename.endswith(".docx")
        or filename.endswith(".doc")
    )

    if not is_pdf and not is_docx:
        raise HTTPException(
            status_code=400, detail="Only PDF and DOCX files are supported"
        )

    if is_pdf:
        import base64

        b64 = base64.b64encode(content).decode("utf-8")
        message_content = [
            {
                "type": "file",
                "file": {
                    "filename": file.filename or "resume.pdf",
                    "file_data": f"data:application/pdf;base64,{b64}",
                },
            },
            {"type": "text", "text": RESUME_EXTRACT_PROMPT},
        ]
    else:
        from docx import Document
        import io

        doc = Document(io.BytesIO(content))
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        if not doc_text.strip():
            raise HTTPException(
                status_code=422, detail="Could not extract text from the document"
            )
        message_content = (
            f"Resume content:\n\n{doc_text}\n\n---\n\n{RESUME_EXTRACT_PROMPT}"
        )

    client = get_openrouter_client()
    completion = client.chat.completions.create(
        model=settings.GEMMA_3N,
        messages=[{"role": "user", "content": message_content}],
    )

    raw = completion.choices[0].message.content or ""
    try:
        parsed = _clean_json_response(raw)
    except json.JSONDecodeError:
        raise HTTPException(
            status_code=500,
            detail="Failed to parse model response as JSON",
        )

    return {
        "candidate_name": parsed.get("candidate_name", ""),
        "candidate_email": parsed.get("candidate_email", ""),
        "skills": parsed.get("skills", []) if isinstance(parsed.get("skills"), list) else [],
        "experience_summary": parsed.get("experience_summary", ""),
        "education": parsed.get("education", ""),
        "years_of_experience": parsed.get("years_of_experience", ""),
        "current_role": parsed.get("current_role", ""),
        "location": parsed.get("location", ""),
        "degree": parsed.get("degree", ""),
        "college": parsed.get("college", ""),
    }
